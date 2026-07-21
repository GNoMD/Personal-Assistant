# -*- coding: utf-8 -*-
"""Convert 一人份一周豆浆早餐配方.md → .docx

目录用 Word 域代码 HYPERLINK \\l \"书签\"，比 w:hyperlink/@w:anchor 更兼容。
"""
from __future__ import annotations

import re
from pathlib import Path

import markdown
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph
from htmldocx import HtmlToDocx

ROOT = Path(__file__).resolve().parents[1]
MD_PATH = ROOT / "md" / "一人份一周豆浆早餐配方.md"
DOCX_PATH = ROOT / "md" / "一人份一周豆浆早餐配方.docx"

SECTION_BOOKMARKS = [
    ("0. 本周要顾全的五件事", "sec_0"),
    ("1. 饮用与制作", "sec_1"),
    ("2. 一人份红线", "sec_2"),
    ("3. 一周配方", "sec_3"),
    ("4. 目标覆盖一览表", "sec_4"),
    ("5. 一周食材用量", "sec_5"),
    ("6. 全套饮用须知", "sec_6"),
    ("7. 每天健身时，豆浆怎么调", "sec_7"),
    ("8. 增肌减脂", "sec_8"),
    ("9. 应急替换表", "sec_9"),
    ("10. 一周水果搭配", "sec_10"),
    ("11. 食材轮换说明", "sec_11"),
    ("13. 练后轻夜宵", "sec_13"),
    ("14. 简要结论", "sec_14"),
]

TOC_LABELS = [
    ("0. 本周要顾全的五件事", "sec_0"),
    ("1. 饮用与制作", "sec_1"),
    ("2. 一人份红线", "sec_2"),
    ("3. 一周配方", "sec_3"),
    ("4. 目标覆盖一览表", "sec_4"),
    ("5. 一周食材用量 + 保存", "sec_5"),
    ("6. 全套饮用须知", "sec_6"),
    ("7. 每天健身加强版", "sec_7"),
    ("8. 增肌减脂：豆浆以外", "sec_8"),
    ("9. 应急替换表", "sec_9"),
    ("10. 一周水果搭配", "sec_10"),
    ("11. 食材轮换说明", "sec_11"),
    ("13. 练后轻夜宵", "sec_13"),
    ("14. 简要结论", "sec_14"),
]


# Compact typography
BODY_PT = 10.0
TABLE_PT = 9.0
TOC_PT = 9.5
H1_PT = 16.0
H2_PT = 12.0
H3_PT = 10.5
ACCENT = RGBColor(0x1A, 0x5C, 0x6B)  # teal-slate
HEADER_FILL = "E8F4F3"
QUOTE_COLOR = RGBColor(0x3D, 0x55, 0x66)


def set_run_font(
    run,
    east_asia: str = "微软雅黑",
    ascii_font: str = "Calibri",
    size_pt: float | None = None,
    *,
    bold: bool | None = None,
    color: RGBColor | None = None,
):
    run.font.name = ascii_font
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), east_asia)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_cell_margins(cell, top=40, bottom=40, left=60, right=60) -> None:
    """Cell margins in twips (1pt ≈ 20 twips). Compact padding."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn("w:tcMar"))
    if tcMar is not None:
        tcPr.remove(tcMar)
    tcMar = OxmlElement("w:tcMar")
    for edge, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def shade_cell(cell, fill_hex: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(BODY_PT)
    normal.font.color.rgb = RGBColor(0x22, 0x2B, 0x35)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(4)
    pf.line_spacing = 1.15

    heading_cfg = {
        1: (H1_PT, Pt(8), Pt(6), ACCENT),
        2: (H2_PT, Pt(10), Pt(4), ACCENT),
        3: (H3_PT, Pt(7), Pt(3), RGBColor(0x2C, 0x3E, 0x50)),
    }
    for level, (size, before, after, color) in heading_cfg.items():
        style = styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.paragraph_format.space_before = before
        style.paragraph_format.space_after = after
        style.paragraph_format.line_spacing = 1.1


def polish_tables(doc: Document) -> None:
    for table in doc.tables:
        table.autofit = True
        for i, row in enumerate(table.rows):
            for cell in row.cells:
                set_cell_margins(cell)
                if i == 0:
                    shade_cell(cell, HEADER_FILL)
                for p in cell.paragraphs:
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.line_spacing = 1.1
                    for run in p.runs:
                        set_run_font(
                            run,
                            size_pt=TABLE_PT,
                            bold=True if i == 0 else None,
                            color=ACCENT if i == 0 else RGBColor(0x22, 0x2B, 0x35),
                        )


def polish_paragraphs(doc: Document) -> None:
    for p in doc.paragraphs:
        style_name = p.style.name if p.style else ""
        text = (p.text or "").strip()

        if style_name == "Heading 1":
            for run in p.runs:
                set_run_font(run, size_pt=H1_PT, bold=True, color=ACCENT)
            continue
        if style_name == "Heading 2":
            for run in p.runs:
                set_run_font(run, size_pt=H2_PT, bold=True, color=ACCENT)
            # subtle bottom spacing already from style
            continue
        if style_name == "Heading 3":
            for run in p.runs:
                set_run_font(run, size_pt=H3_PT, bold=True, color=RGBColor(0x2C, 0x3E, 0x50))
            continue

        # Compact lists
        if style_name.startswith("List"):
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.1
            for run in p.runs:
                set_run_font(run, size_pt=BODY_PT)
            continue

        # Blockquote-like indented paras
        if p.paragraph_format.left_indent and p.paragraph_format.left_indent.pt > 0:
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.1
            for run in p.runs:
                set_run_font(run, size_pt=9.5, color=QUOTE_COLOR)
            continue

        # Empty / separator lines — collapse
        if not text:
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            continue

        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        for run in p.runs:
            set_run_font(run, size_pt=BODY_PT, color=RGBColor(0x22, 0x2B, 0x35))


def md_to_html(text: str) -> str:
    return markdown.markdown(
        text,
        extensions=["tables", "fenced_code", "sane_lists", "nl2br", "smarty"],
        output_format="html5",
    )


def strip_broken_html_anchors(html: str) -> str:
    return re.sub(
        r'<a\s+[^>]*href="#sec-\d+"[^>]*>(.*?)</a>',
        r"\1",
        html,
        flags=re.I | re.S,
    )


def _ensure_pPr(paragraph) -> OxmlElement:
    p = paragraph._p
    pPr = p.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p.insert(0, pPr)
    return pPr


def add_bookmark(paragraph, name: str, bookmark_id: int) -> None:
    """Insert bookmark after pPr so Word recognizes it reliably."""
    p = paragraph._p
    pPr = _ensure_pPr(paragraph)

    # Remove any existing bookmark with same name on this paragraph
    for tag in (qn("w:bookmarkStart"), qn("w:bookmarkEnd")):
        for el in list(p.findall(tag)):
            p.remove(el)

    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))

    # Place start immediately after pPr; end at paragraph end
    pPr.addnext(start)
    p.append(end)


def clear_paragraph(paragraph) -> None:
    p = paragraph._p
    for child in list(p):
        if child.tag == qn("w:pPr"):
            continue
        p.remove(child)


def _make_run_with_text(text: str, *, hyperlink_style: bool = False) -> OxmlElement:
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    if hyperlink_style:
        rStyle = OxmlElement("w:rStyle")
        rStyle.set(qn("w:val"), "Hyperlink")
        rPr.append(rStyle)
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "0563C1")
        rPr.append(color)
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Calibri")
    rFonts.set(qn("w:hAnsi"), "Calibri")
    rFonts.set(qn("w:eastAsia"), "微软雅黑")
    rPr.append(rFonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(TOC_PT * 2)))  # half-points
    rPr.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), str(int(TOC_PT * 2)))
    rPr.append(szCs)
    run.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    run.append(t)
    return run


def add_hyperlink_field(paragraph, text: str, bookmark_name: str) -> None:
    """Word field: HYPERLINK \\l \"bookmark\" — works in desktop Word / WPS."""
    p = paragraph._p

    def append_fld_char(char_type: str) -> None:
        run = OxmlElement("w:r")
        fld = OxmlElement("w:fldChar")
        fld.set(qn("w:fldCharType"), char_type)
        run.append(fld)
        p.append(run)

    append_fld_char("begin")

    instr_run = OxmlElement("w:r")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    # Leading/trailing spaces matter for field parsing
    instr.text = f' HYPERLINK \\l "{bookmark_name}" '
    instr_run.append(instr)
    p.append(instr_run)

    append_fld_char("separate")
    p.append(_make_run_with_text(text, hyperlink_style=True))
    append_fld_char("end")


def bookmark_headings(doc: Document) -> dict[str, str]:
    found: dict[str, str] = {}
    bid = 100  # avoid colliding with any auto ids from htmldocx
    for p in doc.paragraphs:
        if not p.style or p.style.name != "Heading 2":
            continue
        text = (p.text or "").strip()
        for prefix, bname in SECTION_BOOKMARKS:
            if text.startswith(prefix):
                add_bookmark(p, bname, bid)
                found[bname] = text
                bid += 1
                break
    return found


def rebuild_toc(doc: Document, bookmarks: dict[str, str]) -> None:
    paras = list(doc.paragraphs)
    toc_idx = None
    for i, p in enumerate(paras):
        if p.style and p.style.name == "Heading 2" and (p.text or "").strip() == "目录":
            toc_idx = i
            break
    if toc_idx is None:
        raise RuntimeError("未找到「目录」标题")

    toc_paras: list[Paragraph] = []
    for j in range(toc_idx + 1, len(paras)):
        p = paras[j]
        if p.style and p.style.name.startswith("Heading"):
            break
        toc_paras.append(p)

    for p in toc_paras:
        clear_paragraph(p)

    needed = len(TOC_LABELS)
    while len(toc_paras) < needed:
        ref = toc_paras[-1]._p if toc_paras else paras[toc_idx]._p
        new_p = OxmlElement("w:p")
        ref.addnext(new_p)
        toc_paras.append(Paragraph(new_p, paras[toc_idx]._parent))

    for extra in toc_paras[needed:]:
        parent = extra._p.getparent()
        if parent is not None:
            parent.remove(extra._p)
    toc_paras = toc_paras[:needed]

    tip = paras[toc_idx]._p
    # Insert a short hint paragraph right after 目录 heading if useful
    # (skip — keep clean)

    missing = []
    for p, (label, bname) in zip(toc_paras, TOC_LABELS):
        clear_paragraph(p)
        p.style = doc.styles["Normal"]
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.1
        if bname not in bookmarks:
            missing.append(f"{label} -> {bname}")
            run = p.add_run(label)
            set_run_font(run, size_pt=TOC_PT)
            continue
        add_hyperlink_field(p, label, bname)

    if missing:
        print("WARN missing bookmarks:", missing)


def try_update_fields_via_word(docx_path: Path) -> bool:
    """If Microsoft Word is installed, open/save once so HYPERLINK fields resolve."""
    try:
        import win32com.client  # type: ignore
    except ImportError:
        return False

    word = None
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        doc = word.Documents.Open(str(docx_path.resolve()))
        # Update all fields so links become live
        doc.Fields.Update()
        doc.Save()
        doc.Close(False)
        return True
    except Exception as e:
        print(f"Word COM update skipped: {e}")
        return False
    finally:
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass


def main() -> None:
    md = MD_PATH.read_text(encoding="utf-8")
    html = strip_broken_html_anchors(md_to_html(md))

    doc = Document()
    style_document(doc)

    parser = HtmlToDocx()
    parser.table_style = "Table Grid"
    parser.add_html_to_document(html, doc)

    polish_tables(doc)
    polish_paragraphs(doc)

    bookmarks = bookmark_headings(doc)
    rebuild_toc(doc, bookmarks)

    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    out = DOCX_PATH
    try:
        doc.save(str(out))
    except PermissionError:
        out = DOCX_PATH.with_name(DOCX_PATH.stem + "-可跳转目录.docx")
        doc.save(str(out))
        print(f"原文件被占用，已另存: {out}")

    updated = try_update_fields_via_word(out)
    print(f"ok: {out}")
    print(f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} bookmarks={len(bookmarks)} word_fields_updated={updated}")


if __name__ == "__main__":
    main()

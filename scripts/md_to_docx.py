# -*- coding: utf-8 -*-
"""Convert markdown lecture script to Word."""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

MD = Path(r"d:\home\guihua\儿童超声新技术临床授课讲稿.md")
OUT = Path(r"d:\home\guihua\儿童超声新技术临床授课讲稿.docx")


def set_font(run, name="微软雅黑", size=12, bold=False, color=None, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def add_rich_para(doc, text, size=12, bold=False, align=None, space_after=6, indent=0, color=None, italic=False, bullet=False, number=False):
    if bullet:
        p = doc.add_paragraph(style="List Bullet")
    elif number:
        p = doc.add_paragraph(style="List Number")
    else:
        p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.4
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    clean = re.sub(r"\*\*([^*]+)\*\*", r"\1", text) if bullet or number else text
    if bullet or number:
        clean = re.sub(r"^[-*]\s*", "", clean)
        clean = re.sub(r"^\d+\.\s*", "", clean)
    if bullet or number:
        parts = clean.split("**") if "**" in text else [clean]
    else:
        parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        b = bold or (not bullet and not number and i % 2 == 1)
        run = p.add_run(part)
        set_font(run, size=size, bold=b, color=color, italic=italic)
    return p


def parse_table(lines, idx):
    rows = []
    i = idx
    while i < len(lines) and lines[i].strip().startswith("|"):
        cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        if not all(set(c) <= set("-: ") for c in cells):
            clean = [re.sub(r"\*\*([^*]+)\*\*", r"\1", c) for c in cells]
            rows.append(clean)
        i += 1
    return rows, i


def add_table(doc, rows):
    if not rows:
        return
    cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci in range(cols):
            val = row[ci] if ci < len(row) else ""
            table.rows[ri].cells[ci].text = val
            for p in table.rows[ri].cells[ci].paragraphs:
                for run in p.runs:
                    set_font(run, size=10, bold=(ri == 0))
    doc.add_paragraph()


def main():
    lines = MD.read_text(encoding="utf-8").splitlines()
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(2.8)
        sec.right_margin = Cm(2.8)

    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    style.font.size = Pt(12)

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if not line.strip():
            i += 1
            continue

        if line.startswith("# ") and not line.startswith("## "):
            p = doc.add_heading(line[2:].strip(), level=0)
            for run in p.runs:
                set_font(run, size=22, bold=True, color=RGBColor(0x1A, 0x56, 0x8E))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        if line.startswith("## "):
            add_heading = doc.add_heading(line[3:].strip(), level=1)
            for run in add_heading.runs:
                set_font(run, size=16, bold=True, color=RGBColor(0x2E, 0x74, 0xB5))
            i += 1
            continue

        if line.startswith("### "):
            add_heading = doc.add_heading(line[4:].strip(), level=2)
            for run in add_heading.runs:
                set_font(run, size=13, bold=True, color=RGBColor(0x40, 0x40, 0x40))
            i += 1
            continue

        if line.startswith("> "):
            quotes = []
            while i < len(lines) and lines[i].startswith("> "):
                quotes.append(lines[i][2:].strip())
                i += 1
            add_rich_para(doc, "\n".join(quotes), size=11, indent=0.5, color=RGBColor(0x55, 0x55, 0x55), italic=True, space_after=8)
            continue

        if line.strip() == "---":
            doc.add_paragraph()
            i += 1
            continue

        if line.strip().startswith("|"):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue

        if line.startswith("- "):
            add_rich_para(doc, line[2:], bullet=True, space_after=3)
            i += 1
            continue

        if re.match(r"^\d+\.\s", line):
            add_rich_para(doc, line, number=True, space_after=3)
            i += 1
            continue

        if line.startswith("> **") or (line.startswith(">") and i > 0):
            pass

        add_rich_para(doc, line)
        i += 1

    doc.save(str(OUT))
    print("OK:", OUT)


if __name__ == "__main__":
    main()

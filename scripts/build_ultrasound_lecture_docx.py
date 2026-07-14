# -*- coding: utf-8 -*-
"""Generate pediatric ultrasound lecture Word document."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"d:\home\guihua\儿童超声新技术临床宣讲讲稿（Word版）.docx"


def set_font(run, name="微软雅黑", size=12, bold=False, color=None, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def add_para(doc, text, size=12, bold=False, align=None, space_after=6, indent=0, color=None, italic=False):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.4
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        run = p.add_run(part)
        set_font(run, size=size, bold=(bold or i % 2 == 1), color=color, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        colors = {0: RGBColor(0x1A, 0x56, 0x8E), 1: RGBColor(0x2E, 0x74, 0xB5), 2: RGBColor(0x40, 0x40, 0x40)}
        set_font(run, size={0: 22, 1: 16, 2: 14}.get(level, 12), bold=True, color=colors.get(level))
    return p


def add_highlight_box(doc, title, lines):
    """Add a shaded paragraph block."""
    add_para(doc, title, size=11, bold=True, color=RGBColor(0x1A, 0x56, 0x8E))
    for line in lines:
        add_para(doc, "▎ " + line, size=11, indent=0.3, space_after=3)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                set_font(run, size=10, bold=True)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = val
            for p in table.rows[ri + 1].cells[ci].paragraphs:
                for run in p.runs:
                    set_font(run, size=10)
    doc.add_paragraph()


def main():
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(2.8)
        sec.right_margin = Cm(2.8)

    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    style.font.size = Pt(12)

    # ===== 封面信息 =====
    add_para(doc, "复旦大学附属儿科医院厦门医院 · 厦门市儿童医院", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "超声科 · 新技术新项目临床推广宣讲讲稿", size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    add_para(doc, "主讲：陈泽坤、陈晓康  |  受众：临床医生  |  建议时长：15～18分钟", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x66, 0x66, 0x66))
    add_para(doc, "【使用说明】加粗内容为授课重点；灰色框为案例；末尾附申请对接流程。所有病例均已脱敏。", size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x88, 0x88, 0x88), space_after=18)

    # ===== 第一章 开场 =====
    add_heading(doc, "第一章  开场：给临床老师的一封「超声邀请函」", 1)

    add_para(doc, "（互动开场——请环顾四周）", size=11, italic=True, color=RGBColor(0x88, 0x88, 0x88))
    add_para(doc, "各位临床老师，先问三个问题：")
    add_para(doc, "① 有没有遇到过——娃明明该做检查，一听说要进影像科，家长当场开始「情绪管理」？")
    add_para(doc, "② 有没有收到过——「腺样体到底多大？能不能给个准数？」而您只能回复「建议耳鼻喉科再看看」？")
    add_para(doc, "③ 有没有纠结过——危重患儿要置鼻肠管，转运、拍片、全麻，哪一步都像在走钢丝？")
    add_para(doc, "（停顿）如果至少中了一条，今天这 15 分钟，大概率能帮您少打几个「超声科能不能……」的咨询电话。", space_after=10)

    add_para(doc, "我是超声科陈泽坤/陈晓康。今天不讲「机器多贵、参数多牛」，只讲一件事：**13 项新技术，怎么落到您的日常诊疗里，让娃少受罪、让您少纠结。**", bold=False)
    add_para(doc, "三个关键词请记一下：**安全（少辐射少折腾）、精准（能定量能随访）、可协作（MDT 一起上）**。", space_after=12)

    add_heading(doc, "【本章重点】", 2)
    add_highlight_box(doc, "", [
        "超声科定位：从「拍张照」升级为「诊断 + 引导 + 治疗 + 随访」",
        "今日目标：让临床老师知道「什么时候该找我们、找我们能解决什么」",
    ])

    # ===== 第二章 总览 =====
    add_heading(doc, "第二章  新技术全景图：13 项硬菜一张表", 1)
    add_para(doc, "别被 13 个名字吓到——按临床场景分，其实就五类：", space_after=6)

    add_table(doc, ["类别", "技术项目", "临床最常用场景"], [
        ["呼吸/过敏", "ILIT、经颏下扁桃体/腺样体超声", "AR 脱敏、OSA 评估、术后随访"],
        ["消化/营养/阑尾", "CeVUS、CEUS 鼻空肠管、ERAT、胃肠造影", "VUR、肠内营养、非复杂性阑尾炎"],
        ["肌骨/介入", "肌骨神经超声、介入超声 MDT", "斜颈、运动损伤、穿刺活检/引流"],
        ["定量评估", "剪切波弹性、VOCAL 胆囊/巨舌、食管超声", "肝纤维化、BA、BWS、食管病变"],
        ["新生儿/神经", "多模态脊髓脊柱超声", "脊髓拴系、脊柱裂筛查、腰穿定位"],
    ])

    add_para(doc, "（幽默一句）以前超声科像「拍照馆」，现在更像「儿科诊疗外挂」——您负责决策，我们负责把看不见的地方，尽量变成看得见、量得出、跟得上。", space_after=12)

    # ===== 第三章 案例 =====
    add_heading(doc, "第三章  三个真实场景：新技术到底帮了谁", 1)
    add_para(doc, "以下病例均来自我院临床工作场景，**已做脱敏处理**（姓名、住院号、精确日期均已隐去），仅供教学交流。", size=11, color=RGBColor(0x66, 0x66, 0x66), space_after=10)

    add_heading(doc, "案例一  「睡觉像装修的小明」——经颏下超声评估腺样体/扁桃体", 2)
    add_highlight_box(doc, "【脱敏病例摘要】", [
        "患儿男，5 岁，因夜间打鼾、张口呼吸 8 个月就诊。",
        "家长拒绝鼻内镜：「孩子上次在耳鼻喉科哭到缺氧，再也不做了。」",
        "经颏下超声：扁桃体阻塞气道占比 62%（3 度）；腺样体 A/N=0.76（病理性肥大）。",
        "与传统方案对比：无需镇静、无辐射；当日出量化报告，直接支撑 OSA 评估与手术指征讨论。",
        "后续：腺样体切除术后 3 个月超声随访，A/N 降至 0.54，症状明显改善。",
    ])
    add_para(doc, "**核心优势**：把「看起来很大」变成「阻塞占比 51%～75%」「A/N≥0.71」，沟通成本直线下降。探头 5～12 MHz，舌骨上区经皮扫查，娃配合度明显高于鼻内镜。", space_after=10)

    add_heading(doc, "案例二  「PICU 里的营养难题」——CEUS 引导鼻空肠管置入", 2)
    add_highlight_box(doc, "【脱敏病例摘要】", [
        "患儿女，3 岁，重症肺炎合并胃潴留，鼻胃管喂养反复反流，误吸风险高。",
        "需启动鼻空肠营养，但转运至放射科透视置管风险大，家长强烈拒绝多次搬动。",
        "床旁 CEUS 引导：实时观察胃腔→幽门→十二指肠→空肠近段，确认末端位于空肠起始段。",
        "对比数据：徒手盲插成功率 52.2%～76.0%；CEUS 床旁引导——安全、可重复、减少射线与搬动。",
        "结果：一次置管成功，启动肠内营养，未发生误入气管事件。",
    ])
    add_para(doc, "**核心优势**：危重娃「能不动就不动」。胃镜置管成功率虽高（96.2%～100%），但需全麻与转运；**CEUS 是床旁可视化方案，尤其适合 PICU/NICU 场景**。", space_after=10)

    add_heading(doc, "案例三  「黄疸娃的肝硬度问号」——实时剪切波弹性成像", 2)
    add_highlight_box(doc, "【脱敏病例摘要】", [
        "婴儿男，58 天，持续性黄疸、胆汁淤积待查，临床高度怀疑胆道闭锁（BA）。",
        "家长焦虑：「能不能少做有创的？」",
        "剪切波弹性成像定量肝硬度，并结合 VOCAL 评估胆囊收缩：脂餐后 2 小时胆囊收缩 <1/3（收缩功能差），胆囊形态异常。",
        "与穿刺活检相比：弹性成像为**非侵入性、可重复**的肝纤维化评估手段，可用于 BA 术前评估及术后随访。",
        "注：胆囊收缩异常仅为 BA 支持征象，**单项指标不能单独诊断 BA**，需结合综合临床与手术/病理。",
    ])
    add_para(doc, "**核心优势**：把「可能需要反复穿刺」变成「可序列化随访的定量指标」，减轻家庭负担，也为肝移植术后评估提供经济安全的路径。", space_after=12)

    add_heading(doc, "【本章重点】", 2)
    add_highlight_box(doc, "", [
        "案例共同逻辑：少折腾、能定量、能随访",
        "临床医生可直接引用分度/比值/成功率数据与家长及团队沟通",
    ])

    # ===== 第四章 分技术精讲 =====
    add_heading(doc, "第四章  分技术精讲：临床医生速查手册", 1)

    sections = [
        ("4.1  超声引导下淋巴结内免疫治疗（ILIT）", [
            "适用：5 岁以上顽固性 AR，点刺/sIgE 阳性，中重度症状影响睡眠/日常，可伴结膜炎或哮喘。",
            "亮点：疗程约 2 个月 vs 传统脱敏 3～5 年；超声引导避开血管神经，实时观察药液扩散。",
            "找我们：过敏/鼻炎疗效不佳、脱敏依从性差的患儿。",
        ]),
        ("4.2  经颏下超声评估扁桃体/腺样体", [
            "适用：打鼾憋气、抗拒鼻内镜、术后随访、排除其他 OSA 病因。",
            "量化：扁桃体 0～4 度（阻塞占比）；腺样体 A/N 比值（0.50～0.60 正常，≥0.71 病理性，≥0.80 显著）。",
            "找我们：需要「给家长一个数」的 OSA 评估。",
        ]),
        ("4.3  小儿肌骨神经超声", [
            "适用：斜颈、DDH、扳机指、运动损伤、骨骺炎、神经卡压/损伤、软组织肿块等。",
            "找我们：X 线/CT 看不清软组织，或需要动态评估疼痛功能时。",
        ]),
        ("4.4  排泄性尿路超声造影（CeVUS）", [
            "适用：泌尿道感染、肾盂输尿管积水、产前高级别积水、VUR 随访等。",
            "亮点：无辐射、实时动态；VUR I～V 级评估；禁忌包括急性感染、造影剂过敏等。",
            "找我们：需要反复评估反流的泌尿外科/肾内科合作病例。",
        ]),
        ("4.5  床旁 CEUS 引导鼻空肠管置入", [
            "适用：鼻饲不耐受、胃潴留、误吸高风险、需肠内营养支持。",
            "关键：幽门通过是置管成功关键；末端确认于空肠起始段。",
            "找我们：PICU/NICU/消化科置管困难、需减少搬动与射线的病例。",
        ]),
        ("4.6  超声造影引导 ERAT（保阑尾治疗）", [
            "适用：急性非复杂性阑尾炎，希望避免开刀、保留阑尾功能者。",
            "亮点：门诊可开展、无创无疤；超声造影评估清洗排尽及穿孔风险。",
            "找我们：外科/急诊评估可保守治疗的阑尾炎病例。",
        ]),
        ("4.7  胃肠超声造影", [
            "适用：罗马 IV 功能性消化不良、可疑 GER、胃十二指肠病变。",
            "亮点：非侵入、无辐射，评估胃动力与胃排空，优于核素的部分场景。",
        ]),
        ("4.8  介入超声 + MDT", [
            "项目：淋巴管瘤硬化、PICC/鼻空肠置管、肝/肾/浅表肿物活检、脓肿引流等。",
            "亮点：较盲穿损伤小、成功率高；需超声+临床+病理协作。",
        ]),
        ("4.9  实时剪切波弹性成像（肝纤维化）", [
            "适用：新生儿黄疸、胆汁淤积、BA 术前术后、慢性肝病、肝移植随访。",
            "亮点：定量肝硬度，非侵入、可重复。",
        ]),
        ("4.10  三维 VOCAL——胆囊收缩功能", [
            "评判（餐后 2 h）：缩小 >2/3 良好；<1/2 可疑；<1/3 差；同空腹为无收缩。",
            "正常脂餐后 45～60 min 体积减小约 60%。BA 支持征象但非单独诊断依据。",
        ]),
        ("4.11  三维 VOCAL——巨舌症舌体容积", [
            "适用：BWS 等巨舌症术前评估及术后容积随访。",
        ]),
        ("4.12  小儿食管超声", [
            "适用：先天性闭锁/瘘管、术后狭窄、GER、食管炎、异物、支架引导等。",
        ]),
        ("4.13  多模态超声——新生儿脊髓/脊柱", [
            "适用：脊髓拴系、脊柱裂、脊膜膨出、椎管内肿瘤等。",
            "要点：足月儿圆锥 L2/L3 及以上；终丝止于 S2；可引导腰穿、动态随访。",
            "禁忌：病情极重体位受限、局部皮肤感染等。",
        ]),
    ]

    for title, bullets in sections:
        add_heading(doc, title, 2)
        for b in bullets:
            add_para(doc, "• " + b, space_after=3)
        add_para(doc, "", space_after=6)

    # ===== 第五章 幽默缓冲 =====
    add_heading(doc, "第五章  超声科「辟谣时间」——三句大实话", 1)
    add_para(doc, "① **超声不是万能的**——比如 BA 不能只看胆囊收缩一个指标，我们比您更怕「一锤定音」的误解。", space_after=4)
    add_para(doc, "② **超声不是越贵越好**——关键是「娃配合度 + 信息增量 + 能不能改变决策」。", space_after=4)
    add_para(doc, "③ **找我们不是「加项」**——很多时候是帮您少做一个有辐射/有麻醉的检查，这是减负，不是加码。", space_after=12)

    # ===== 第六章 落地指引 =====
    add_heading(doc, "第六章  落地指引：临床老师如何对接开展", 1)

    add_heading(doc, "6.1  什么时候该开单/会诊超声科？", 2)
    add_table(doc, ["您的场景", "建议对接项目", "对接方式"], [
        ["打鼾/OSA 纠结，娃不配合鼻内镜", "经颏下扁桃体/腺样体超声", "门诊超声申请 + 注明 OSA 评估"],
        ["危重娃需鼻空肠营养", "CEUS 引导置管", "床旁会诊 / PICU 联络超声"],
        ["反复 UTIs/VUR 随访", "CeVUS", "泌尿外科联合申请"],
        ["BA/胆汁淤积需肝硬度随访", "剪切波弹性 + VOCAL 胆囊", "消化/新生儿科联合评估"],
        ["新生儿腰穿/脊髓问题", "多模态脊髓脊柱超声", "新生儿科床旁会诊"],
    ])

    add_heading(doc, "6.2  申请与对接流程（建议路径）", 2)
    add_para(doc, "**Step 1  临床评估**——确认符合各技术适应证/禁忌证，向家属说明检查目的（宣教口径见 PPT）。", space_after=4)
    add_para(doc, "**Step 2  提交申请**——HIS 开具超声检查项目，**备注栏注明「新技术项目 + 具体名称」**（如：经颏下腺样体超声、CEUS 置管等），便于分诊。", space_after=4)
    add_para(doc, "**Step 3  超声科分诊**——门诊至超声科服务台；床旁项目由科室电话联络超声科协调（PICU/NICU/急诊优先）。", space_after=4)
    add_para(doc, "**Step 4  报告与 MDT**——量化报告（分度、A/N、弹性值等）回传临床；复杂病例可发起 **MDT 讨论**（超声 + 临床 + 病理/麻醉/营养）。", space_after=4)
    add_para(doc, "**Step 5  随访**——术后/术后系列评估请同一模板复查，便于纵向对比（尤其 OSA、BA、巨舌术后）。", space_after=10)

    add_heading(doc, "6.3  配套支持", 2)
    add_highlight_box(doc, "", [
        "授课/跟诊：欢迎各科室预约「新技术跟诊半天」，超声医生现场讲解图像与报告。",
        "病例讨论：提供脱敏典型图像与报告模板，支持科室内部培训。",
        "MDT 协作：介入、ERAT、置管、活检类项目需提前沟通患儿状态与凝血/感染情况。",
        "联系方式：请通过院内超声科科室电话 / 钉钉工作群联络（以医院最新通讯录为准）。",
    ])

    # ===== 结尾 =====
    add_heading(doc, "结语  三句话带走", 1)
    add_para(doc, "**技术革新，安全先行；超声赋能，精准诊疗；多模态融合，多学科配合。**", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    add_para(doc, "各位临床老师，娃的检查体验，一半在您怎么选检查，一半在我们怎么把检查做好。欢迎把今天讲到的 13 项新技术，写进您的诊疗路径里——**需要的时候，超声科随叫随到。**", space_after=8)
    add_para(doc, "谢谢大家！欢迎提问。", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    add_heading(doc, "【附录】讲师备注与合规提示", 2)
    add_table(doc, ["项目", "说明"], [
        ["时长", "全文约 3500 字，建议 15～18 分钟；案例部分可扩展至 20 分钟"],
        ["互动", "开场三问、案例后提问「您科室最常见哪一类？」"],
        ["合规", "治疗类项目须强调「需评估适应证与禁忌证」；案例已脱敏；避免绝对化疗效承诺"],
        ["PPT 配合", "案例对应 PPT 第 7～13、24～26、41～51、60～67 页"],
        ["数据引用", "置管成功率、A/N 比值、VUR 分级、胆囊收缩标准均来自科室 PPT 原文"],
    ])

    doc.save(OUT)
    print("Saved:", OUT)


if __name__ == "__main__":
    main()
